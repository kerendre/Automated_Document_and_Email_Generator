import data_cleaning_functions
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx
import adding_objects_file
import send_mail_tab
import pandas as pd

# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
# Define variables for pictures and their descriptions
pic1 = "pic1.JPG"
pic1_description = "pic1 is description of pic1"

pic2 = "pic2.PNG"
pic2_description = "pic2_description"

# Define variables for table1
table1_origin_location_n_name = "DATABASES/input_xlsx_file.xlsx"  # tourism_israel.xlsx
table1_sheet_name = "Sheet1"
table1_number_of_rows_2_skip = []
subtitle_to_the_table1 = "This is subtitle to  the table "

# Define variables for table2
table2_origin_location_n_name = "DATABASES/Contact_Info.xlsx"
table2_sheet_name = "Sheet1"
table2_number_of_rows_2_skip = []

# Define the column that will be used to merge the dataframes
merge_dfs_on_column = "municipal"

# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
# Choose 1. or 2. but not both

# 1. Open the Word document
# doc = docx.Document("input_word_file.docx")

# 2. Create an instance of a word document
doc = docx.Document()

# add a  header and add hyperlink to the header
adding_objects_file.add_hyperlink_to_header(doc, 'https://www.linkedin.com/in/keren-drevin/', 'Linkdein/Keren Drevin')
adding_objects_file.add_hyperlink_to_header(doc, 'https://github.com/kerendre',
                                            '                                                                              GitHub/Keren Drevin')
adding_objects_file.add_hyperlink_to_header(doc, 'https://www.linkedin.com/in/keren-drevin/',
                                            "\n ________________________________________________________________________________________________________")
# End of 2.
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

# Open the source Word document, the origin of the text in my document
source_doc = docx.Document("DATABASES/input_word_file.docx")

# Loop over all paragraphs in the source document and add them to the destination document
for i, paragraph in enumerate(source_doc.paragraphs):
    if i == 0:  # first paragraph is the title
        title = doc.add_paragraph(paragraph.text) # Add the title paragraph to the destination document
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER # Center align the title paragraph
        title.style = 'Heading 2' # Set the style of the title paragraph to 'Heading 2
        for run in title.runs:
            run.font.color.rgb = docx.shared.RGBColor(127, 0, 0)  # Dark red
            run.font.underline = True  # underline

    else:
        doc.add_paragraph(paragraph.text)

# ~~~~~~~~~~~~~~~~~~~~~~~end 20.02.23~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

# Adding a picture to the Word document
adding_objects_file.Add_pic(doc, pic_name=pic1, pic_description=pic1_description)

# Adding another picture to the Word document
adding_objects_file.Add_pic(doc, pic_name=pic2, pic_description=pic2_description)

# Reading the data from the excel file into a pandas dataframe)
df = pd.read_excel(table1_origin_location_n_name, sheet_name=table1_sheet_name, skiprows=table1_number_of_rows_2_skip)

# Convert the specified column to a string in the desired format "%d/%m/%Y"
df = data_cleaning_functions.format_dates_2_d_m_Y(df, "first_visi")
df = data_cleaning_functions.format_dates_2_d_m_Y(df, "second_vis")
df = data_cleaning_functions.format_dates_2_d_m_Y(df, "third_visi")

# Remove all columns that end with 'heb' or 'heb '
df = df.loc[:, ~df.columns.str.endswith(('heb', 'heb '))]

# Drop columns 'GIS_ID', 'X', and 'Y'
df = df.drop(['GIS_ID', 'X', 'Y'], axis=1)

# Read contact information from another Excel file into a pandas data
df_contacts = pd.read_excel(table2_origin_location_n_name, sheet_name=table2_sheet_name,
                            skiprows=table2_number_of_rows_2_skip)

# Merge data from the two dataframes on a specific column
merged_df = data_cleaning_functions.merge_data(df, df_contacts, merge_dfs_on_column)

# Drop the 'send mail' column from the merged dataframe
merged_df_with_no_send_mail_col = merged_df.drop('send mail', axis=1)


# Adding the dataframe as a table to the Word document
doc = adding_objects_file.Add_table(doc, merged_df_with_no_send_mail_col, subtitle_to_the_table=subtitle_to_the_table1)

# Save the Word document as "doc1.docx"
doc_to_attach = doc.save("doc1.docx")

# ~~~~~~~~~~~~~ "Extract the email text from a file and save it into the 'mail_text' variable."~~~~~~~~~~~~~~~~~~~~
# Open the Word document where the test is saved
input_mail_text_doc = docx.Document('DATABASES/mail_text.docx')

# Extract the text from the document
mail_text = '\n'.join([paragraph.text for paragraph in input_mail_text_doc.paragraphs])
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

# Convert the df_contacts dataframe to a list of dictionaries
dict_list = merged_df.to_dict(orient='records')
for num in range(0, len(dict_list)):
    name = dict_list[num]['name']
    mail = dict_list[num]['mail']
    company_name = dict_list[num]['municipal']

    to_send_a_mail = dict_list[num]['send mail']
    # Only send mail if the value in the 'send mail' column is 'V' or 'v'
    if to_send_a_mail == 'V' or to_send_a_mail == 'v':  # Only send mail if the value in send mail column value is not none.
        send_mail_tab.send_mail(mail, name, company_name, "doc1.docx",mail_text)
