# Import docx NOT python-docx
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def Add_pic(doc, pic_name, pic_description):

    """
    This function adds a picture to a Word document along with a description and spacing.

    Parameters:
    doc (docx.Document): The Word document to which the picture is added.
    pic_name (str): The file name of the picture.
    pic_description (str): The description to be added under the picture.

    Returns:
    None
    """

    # Get the page width
    page_width = doc.sections[0].page_width

    # Add the picture to the Word document
    pic1 = doc.add_picture(f"PICS/{pic_name}")

    # Fit the width of the first picture to the page width
    pic1.width = int(page_width * 0.65)

    # Add a description under the picture pic1
    description_pic = doc.add_paragraph(pic_description)

    # Center the description
    description_pic.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    # Add two lines of space
    doc.add_paragraph("\n\n")

    # Adding a page break
    doc.add_page_break()


def Add_table(doc, df, subtitle_to_the_table):
    """  Add a table with data from a Pandas DataFrame to a Word document.

    Parameters:
    - doc: a Word document object, created using the `docx.Document()` constructor.
    - df: a Pandas DataFrame object containing the data to be added to the table.
    - subtitle_to_the_table: a string representing the subtitle of the table.

    Returns:
    - the Word document object, with the table added.
    """

    # Replace NaN values with an empty string
    df = df.fillna('')

    # ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    # Add a Title to the document,
    # Chose 1. showing all title options or 2. the real title
    # 1.
    # for num in range(0,6):
    #     heading =doc.add_heading(subtitle_to_the_table, num)
    #     heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2.
    # Add a heading to the document with the specified subtitle
    heading = doc.add_heading(subtitle_to_the_table, 2)

    # Center the heading
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add some style to the heading
    for run in heading.runs:
        run.font.color.rgb = docx.shared.RGBColor(127, 0, 0)  # dark red
        run.font.underline = True  # underline
    # ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    # Creating a table object
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

    # Add data to the table
    row = table.rows[0].cells
    for i in range(df.shape[1]):
        row[i].text = df.columns[i]

    for i in range(df.shape[1]):
        row[i].text = df.columns[i]
    for row in range(df.shape[0]):
        if not df.iloc[row].isnull().all():
            row_cells = table.rows[row + 1].cells
            for col in range(df.shape[1]):
                if isinstance(df.iloc[row, col], float):
                    row_cells[col].text = f"{df.iloc[row, col]:.2f}"
                else:
                    row_cells[col].text = str(df.iloc[row, col])

    # Add style to the table
    table.style = 'Colorful List'

    return doc


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def add_hyperlink_to_header(document, url, text):
    """
    A function that adds a hyperlink to the header of a word document.

    :param document: The document to add the hyperlink to.
    :param url: A string containing the required url.
    :param text: The text displayed for the url.
    :return: The header object.
    """
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]

    hyperlink = add_hyperlink(paragraph, url, text, 'blue', True)

    # return header
    return hyperlink
